"""PDF and Word report generation for the truck-impact simulator."""

from __future__ import annotations

import math
import re
from datetime import datetime
from io import BytesIO

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from constants import (
    AIS_DESCRIPTIONS,
    APA_REFERENCES,
    G,
    ROSEN_SANDER_COEFS,
    SURFACE_LABELS,
    SURFACE_MU_R,
)
from physics import SimulationResult


def _strip_html(text: str) -> str:
    return re.sub(r"<[^>]+>", "", text).replace("&amp;", "&").replace("&gt;", ">").replace("&lt;", "<")


def generate_pdf(
    result: SimulationResult,
    slope_deg: float,
    distance_m: float,
    mass_kg: float,
    surface_key: str,
) -> bytes:
    surface_label = SURFACE_LABELS[surface_key]
    mu_r = SURFACE_MU_R[surface_key]
    theta_rad = math.radians(slope_deg)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title="Memoria de Cálculo — Simulador de Impacto",
        author="Simulador de Impacto",
    )

    styles = getSampleStyleSheet()

    s_title = ParagraphStyle("s_title", parent=styles["Title"],
        fontSize=17, textColor=colors.HexColor("#1565c0"), alignment=TA_CENTER, spaceAfter=4)
    s_sub = ParagraphStyle("s_sub", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#555"), alignment=TA_CENTER, spaceAfter=18)
    s_h1 = ParagraphStyle("s_h1", parent=styles["Heading1"],
        fontSize=13, textColor=colors.HexColor("#1565c0"), spaceBefore=14, spaceAfter=5)
    s_h2 = ParagraphStyle("s_h2", parent=styles["Heading2"],
        fontSize=11, textColor=colors.HexColor("#333"), spaceBefore=9, spaceAfter=3)
    s_body = ParagraphStyle("s_body", parent=styles["Normal"],
        fontSize=10, leading=14, alignment=TA_JUSTIFY, spaceAfter=4)
    s_eq = ParagraphStyle("s_eq", parent=styles["Normal"],
        fontSize=10, leading=14, leftIndent=1 * cm, fontName="Courier",
        textColor=colors.HexColor("#1a237e"), spaceAfter=2)
    s_warn = ParagraphStyle("s_warn", parent=styles["Normal"],
        fontSize=10, leading=14, textColor=colors.HexColor("#c62828"),
        backColor=colors.HexColor("#ffebee"), borderColor=colors.HexColor("#c62828"),
        borderWidth=1, borderPad=6, spaceAfter=6)
    s_ok = ParagraphStyle("s_ok", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#2e7d32"), fontName="Helvetica-Bold", spaceAfter=4)
    s_ref = ParagraphStyle("s_ref", parent=styles["Normal"],
        fontSize=9, leading=13, leftIndent=0.8 * cm, firstLineIndent=-0.8 * cm, spaceAfter=5)
    s_footer = ParagraphStyle("s_footer", parent=styles["Normal"],
        fontSize=8, textColor=colors.grey, alignment=TA_CENTER)

    _tbl_header = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1565c0")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#e3f2fd")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ])
    _tbl_results = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e65100")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (1, 0), (2, -1), "RIGHT"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fff3e0")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ])

    story = []

    # ── Header ──────────────────────────────────────────────────────────────
    story += [
        Paragraph("MEMORIA DE CÁLCULO", s_title),
        Paragraph("Simulador de Impacto — Camioneta Rodando por Pendiente", s_sub),
        Paragraph(f"Generado el {datetime.now().strftime('%d de %B de %Y, %H:%M')}", s_sub),
        HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1565c0")),
        Spacer(1, 0.4 * cm),
    ]

    # ── 1. Parámetros de entrada ─────────────────────────────────────────────
    story.append(Paragraph("1. Parámetros de Entrada", s_h1))
    t_input = Table([
        ["Parámetro", "Símbolo", "Valor", "Unidad"],
        ["Ángulo de pendiente", "θ", f"{slope_deg:.1f}", "°"],
        ["Distancia a la persona", "d", f"{distance_m:.0f}", "m"],
        ["Masa de la camioneta", "m", f"{mass_kg:.0f}", "kg"],
        ["Superficie", "—", surface_label, "—"],
        ["Coef. fricción de rodadura", "μ_r", f"{mu_r:.4f}", "adimensional"],
        ["Aceleración gravitacional", "g", f"{G}", "m/s²"],
    ], colWidths=[6 * cm, 2.5 * cm, 3 * cm, 2 * cm])
    t_input.setStyle(_tbl_header)
    story += [t_input, Spacer(1, 0.4 * cm)]

    # ── 2. Modelo físico ─────────────────────────────────────────────────────
    story.append(Paragraph("2. Modelo Físico y Desarrollo del Cálculo", s_h1))
    story.append(Paragraph(
        "Se modela la camioneta como un sólido rígido en movimiento rectilíneo uniformemente "
        "acelerado sobre un plano inclinado, con fricción de rodadura pura (sin deslizamiento), "
        "despreciando resistencia aerodinámica y considerando velocidad inicial nula.",
        s_body))

    # 2.1 Condición de arranque
    story.append(Paragraph("2.1  Condición de arranque", s_h2))
    story.append(Paragraph(
        "Para que la camioneta ruede, la componente gravitacional tangencial debe superar "
        "la fuerza de fricción de rodadura:", s_body))
    story.append(Paragraph("tan θ  >  μ_r", s_eq))
    tan_val = math.tan(theta_rad)
    story.append(Paragraph(
        f"tan({slope_deg:.1f}°) = {tan_val:.4f}  {'>' if result.will_roll else '≤'}  μ_r = {mu_r:.4f}",
        s_eq))
    if result.will_roll:
        story.append(Paragraph("✔  La camioneta ARRANCA y rueda libremente pendiente abajo.", s_ok))
    else:
        story.append(Paragraph(
            "✘  La camioneta NO arranca. La fricción de rodadura supera la componente "
            "gravitacional. No hay impacto.", s_warn))

    # 2.2 Aceleración neta
    story.append(Paragraph("2.2  Aceleración neta a lo largo de la pendiente", s_h2))
    story.append(Paragraph(
        "La Segunda Ley de Newton sobre el eje de la pendiente, considerando solo "
        "la componente peso y la fricción de rodadura:", s_body))
    story += [
        Paragraph("a  =  g · (sin θ  −  μ_r · cos θ)", s_eq),
        Paragraph(f"a  =  {G} · (sin {slope_deg:.1f}° − {mu_r} · cos {slope_deg:.1f}°)", s_eq),
        Paragraph(
            f"a  =  {G} · ({math.sin(theta_rad):.4f} − {mu_r} · {math.cos(theta_rad):.4f})", s_eq),
        Paragraph(
            f"a  =  {G} · ({math.sin(theta_rad):.4f} − {mu_r * math.cos(theta_rad):.4f})", s_eq),
        Paragraph(f"a  =  {result.acceleration:.4f} m/s²", s_eq),
    ]

    if result.will_roll:
        # 2.3 Velocidad de impacto
        story.append(Paragraph("2.3  Velocidad de impacto", s_h2))
        story.append(Paragraph(
            "Con aceleración constante y velocidad inicial nula, la cinemática da:", s_body))
        story += [
            Paragraph("v²  =  v₀²  +  2 · a · d   →   v  =  √(2 · a · d)   (v₀ = 0)", s_eq),
            Paragraph(f"v  =  √(2 · {result.acceleration:.4f}  ·  {distance_m:.1f})", s_eq),
            Paragraph(f"v  =  √({2 * result.acceleration * distance_m:.4f})", s_eq),
            Paragraph(
                f"v  =  {result.v_impact_mps:.4f} m/s  =  "
                f"{result.v_impact_mps:.4f} × 3.6  =  {result.v_impact_kmh:.2f} km/h", s_eq),
        ]

        # 2.4 Energía cinética
        story.append(Paragraph("2.4  Energía cinética en el impacto", s_h2))
        story.append(Paragraph(
            "La energía cinética representa la capacidad de deformación y daño al momento del choque:",
            s_body))
        story += [
            Paragraph("E_k  =  ½ · m · v²", s_eq),
            Paragraph(f"E_k  =  0.5  ·  {mass_kg:.0f}  ·  ({result.v_impact_mps:.4f})²", s_eq),
            Paragraph(
                f"E_k  =  {result.kinetic_energy_j:,.2f} J  "
                f"=  {result.kinetic_energy_j / 1000:.4f} kJ", s_eq),
        ]
        story.append(Paragraph(
            "Nota: la masa NO afecta la velocidad de impacto (se cancela en la ecuación de "
            "cinemática), pero sí determina la energía transferida al peatón.", s_body))

        # 2.5 Probabilidad de fatalidad
        a_f = ROSEN_SANDER_COEFS["a_fatal"]
        b_f = ROSEN_SANDER_COEFS["b_fatal"]
        a_s = ROSEN_SANDER_COEFS["a_serious"]
        b_s = ROSEN_SANDER_COEFS["b_serious"]
        z_f = a_f - b_f * result.v_impact_kmh
        z_s = a_s - b_s * result.v_impact_kmh

        story.append(Paragraph("2.5  Probabilidad de fatalidad del peatón (Rosén & Sander, 2009)", s_h2))
        story.append(Paragraph(
            "Modelo logístico ajustado a datos estadísticos de accidentes reales. "
            "La variable independiente es la velocidad de impacto del vehículo en km/h:", s_body))
        story += [
            Paragraph(f"P_fatal(v)  =  1 / (1 + exp({a_f} − {b_f} · v_kmh))", s_eq),
            Paragraph(
                f"P_fatal  =  1 / (1 + exp({a_f} − {b_f} · {result.v_impact_kmh:.2f}))", s_eq),
            Paragraph(f"P_fatal  =  1 / (1 + e^({z_f:.4f}))  =  {result.fatality_probability*100:.2f}%",
                      s_eq),
            Spacer(1, 0.15 * cm),
            Paragraph(f"P_lesión_grave (AIS3+)  =  1 / (1 + exp({a_s} − {b_s} · v_kmh))", s_eq),
            Paragraph(
                f"P_lesión_grave  =  1 / (1 + e^({z_s:.4f}))  "
                f"=  {result.serious_injury_probability*100:.2f}%", s_eq),
        ]

    # ── 3. Resultados ────────────────────────────────────────────────────────
    story.append(Paragraph("3. Resumen de Resultados", s_h1))

    if not result.will_roll:
        story.append(Paragraph(
            "Con los parámetros ingresados la camioneta NO se desplaza. "
            "No hay impacto y no aplican métricas de severidad.", s_warn))
    else:
        t_res = Table([
            ["Magnitud calculada", "Valor", "Unidad"],
            ["Aceleración neta en la pendiente", f"{result.acceleration:.4f}", "m/s²"],
            ["Velocidad de impacto", f"{result.v_impact_mps:.3f}", "m/s"],
            ["Velocidad de impacto", f"{result.v_impact_kmh:.2f}", "km/h"],
            ["Energía cinética al impacto", f"{result.kinetic_energy_j:,.1f}", "J"],
            ["Energía cinética al impacto", f"{result.kinetic_energy_j / 1000:.3f}", "kJ"],
            ["Probabilidad de fatalidad peatón", f"{result.fatality_probability*100:.1f}", "%"],
            ["Probabilidad de lesión grave (AIS3+)", f"{result.serious_injury_probability*100:.1f}", "%"],
            ["Nivel AIS (escala cualitativa)", result.ais_label, "—"],
        ], colWidths=[8.5 * cm, 3 * cm, 2 * cm])
        t_res.setStyle(_tbl_results)
        story += [t_res, Spacer(1, 0.3 * cm)]

        # AIS clinical interpretation
        for desc in AIS_DESCRIPTIONS:
            ais_key = desc["level"].split("—")[0].strip()
            if result.ais_label.startswith(ais_key):
                story.append(Paragraph("Interpretación clínica (nivel AIS activo):", s_h2))
                ais_color = colors.HexColor(desc["color"])
                for label, key in [
                    ("Rango de velocidad", "range"),
                    ("Descripción clínica", "clinical"),
                    ("Ejemplos de lesiones", "examples"),
                    ("Pronóstico", "prognosis"),
                    ("Contexto", "context"),
                ]:
                    story.append(Paragraph(
                        f"<b>{label}:</b> {_strip_html(desc[key])}", s_body))
                break

    # ── 4. Limitaciones ──────────────────────────────────────────────────────
    story.append(Paragraph("4. Alcance y Limitaciones del Modelo", s_h1))
    story.append(Paragraph(
        "Este modelo es una herramienta <b>didáctica</b>. Las siguientes simplificaciones "
        "deben tenerse en cuenta antes de extrapolar los resultados:", s_body))
    for item in [
        "Se desprecia la resistencia aerodinámica (relevante a velocidades > 60 km/h).",
        "Se asume rodadura pura sin deslizamiento de ruedas.",
        "Los coeficientes μ_r son valores típicos y pueden variar con temperatura, desgaste y carga.",
        "La curva de Rosén & Sander asume impacto frontal de vehículo ligero contra peatón adulto en condiciones europeas.",
        "La escala AIS es cualitativa; la evaluación médica real requiere examen clínico y estudios de imagen.",
        "ESTE DOCUMENTO NO DEBE USARSE PARA ANÁLISIS FORENSE NI DECISIONES DE INGENIERÍA DE SEGURIDAD.",
    ]:
        story.append(Paragraph(f"• {item}", s_body))

    # ── 5. Referencias ───────────────────────────────────────────────────────
    story.append(Paragraph("5. Referencias Bibliográficas (APA 7.ª ed.)", s_h1))
    for ref in APA_REFERENCES:
        story.append(Paragraph(_strip_html(ref), s_ref))

    # Footer
    story += [
        Spacer(1, 0.5 * cm),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#bbbbbb")),
        Paragraph(
            f"Generado automáticamente — Simulador de Impacto · "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
            s_footer,
        ),
    ]

    doc.build(story)
    return buffer.getvalue()


def generate_docx(
    result: SimulationResult,
    slope_deg: float,
    distance_m: float,
    mass_kg: float,
    surface_key: str,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor, Cm

    surface_label = SURFACE_LABELS[surface_key]
    mu_r = SURFACE_MU_R[surface_key]
    theta_rad = math.radians(slope_deg)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    def _heading(text: str, level: int) -> None:
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) if level == 1 else RGBColor(0x33, 0x33, 0x33)

    def _para(text: str, bold: bool = False, italic: bool = False, mono: bool = False,
              color: RGBColor | None = None, indent: float = 0) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        if color:
            run.font.color.rgb = color

    def _table(headers: list[str], rows: list[list[str]], col_widths: list[float]) -> None:
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            from docx.oxml import OxmlElement
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1565C0")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        for ri, row in enumerate(rows):
            for ci, cell_text in enumerate(row):
                t.rows[ri + 1].cells[ci].text = cell_text
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
        doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("MEMORIA DE CÁLCULO")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    sub_p = doc.add_paragraph("Simulador de Impacto — Camioneta Rodando por Pendiente")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.size = Pt(12)

    date_p = doc.add_paragraph(f"Generado el {datetime.now().strftime('%d de %B de %Y, %H:%M')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_paragraph()

    # 1. Input parameters
    _heading("1. Parámetros de Entrada", 1)
    _table(
        ["Parámetro", "Símbolo", "Valor", "Unidad"],
        [
            ["Ángulo de pendiente", "θ", f"{slope_deg:.1f}", "°"],
            ["Distancia a la persona", "d", f"{distance_m:.0f}", "m"],
            ["Masa de la camioneta", "m", f"{mass_kg:.0f}", "kg"],
            ["Superficie", "—", surface_label, "—"],
            ["Coef. fricción de rodadura", "μ_r", f"{mu_r:.4f}", "adimensional"],
            ["Aceleración gravitacional", "g", f"{G}", "m/s²"],
        ],
        [6.5, 2.5, 3.0, 2.5],
    )

    # 2. Physics
    _heading("2. Modelo Físico y Desarrollo del Cálculo", 1)
    _para(
        "Se modela la camioneta como un sólido rígido en movimiento rectilíneo uniformemente "
        "acelerado, con fricción de rodadura pura y velocidad inicial nula."
    )

    _heading("2.1  Condición de arranque", 2)
    _para("tan θ  >  μ_r", mono=True, indent=1.0)
    tan_val = math.tan(theta_rad)
    _para(
        f"tan({slope_deg:.1f}°) = {tan_val:.4f}  {'>' if result.will_roll else '≤'}  μ_r = {mu_r:.4f}",
        mono=True, indent=1.0)
    if result.will_roll:
        _para("✔  La camioneta ARRANCA.", bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
    else:
        _para("✘  La camioneta NO arranca. No hay impacto.", bold=True, color=RGBColor(0xC6, 0x28, 0x28))

    _heading("2.2  Aceleración neta", 2)
    for eq in [
        "a  =  g · (sin θ  −  μ_r · cos θ)",
        f"a  =  {G} · (sin {slope_deg:.1f}° − {mu_r} · cos {slope_deg:.1f}°)",
        f"a  =  {G} · ({math.sin(theta_rad):.4f} − {mu_r * math.cos(theta_rad):.4f})",
        f"a  =  {result.acceleration:.4f} m/s²",
    ]:
        _para(eq, mono=True, indent=1.0)

    if result.will_roll:
        _heading("2.3  Velocidad de impacto", 2)
        for eq in [
            "v  =  √(2 · a · d)    (v₀ = 0)",
            f"v  =  √(2 · {result.acceleration:.4f}  ·  {distance_m:.1f})",
            f"v  =  {result.v_impact_mps:.4f} m/s  =  {result.v_impact_kmh:.2f} km/h",
        ]:
            _para(eq, mono=True, indent=1.0)

        _heading("2.4  Energía cinética", 2)
        for eq in [
            "E_k  =  ½ · m · v²",
            f"E_k  =  0.5 · {mass_kg:.0f} · ({result.v_impact_mps:.4f})²",
            f"E_k  =  {result.kinetic_energy_j:,.2f} J  =  {result.kinetic_energy_j/1000:.4f} kJ",
        ]:
            _para(eq, mono=True, indent=1.0)

        _heading("2.5  Probabilidades (Rosén & Sander, 2009)", 2)
        a_f, b_f = ROSEN_SANDER_COEFS["a_fatal"], ROSEN_SANDER_COEFS["b_fatal"]
        a_s, b_s = ROSEN_SANDER_COEFS["a_serious"], ROSEN_SANDER_COEFS["b_serious"]
        z_f = a_f - b_f * result.v_impact_kmh
        z_s = a_s - b_s * result.v_impact_kmh
        for eq in [
            f"P_fatal  =  1 / (1 + e^({z_f:.4f}))  =  {result.fatality_probability*100:.2f}%",
            f"P_lesion_grave  =  1 / (1 + e^({z_s:.4f}))  =  {result.serious_injury_probability*100:.2f}%",
        ]:
            _para(eq, mono=True, indent=1.0)

    # 3. Results table
    _heading("3. Resumen de Resultados", 1)
    if not result.will_roll:
        _para("La camioneta no se desplaza. No hay impacto.", bold=True,
              color=RGBColor(0xC6, 0x28, 0x28))
    else:
        _table(
            ["Magnitud", "Valor", "Unidad"],
            [
                ["Aceleración neta", f"{result.acceleration:.4f}", "m/s²"],
                ["Velocidad de impacto", f"{result.v_impact_mps:.3f}", "m/s"],
                ["Velocidad de impacto", f"{result.v_impact_kmh:.2f}", "km/h"],
                ["Energía cinética", f"{result.kinetic_energy_j:,.1f}", "J"],
                ["Energía cinética", f"{result.kinetic_energy_j/1000:.3f}", "kJ"],
                ["Prob. fatalidad peatón", f"{result.fatality_probability*100:.1f}", "%"],
                ["Prob. lesión grave (AIS3+)", f"{result.serious_injury_probability*100:.1f}", "%"],
                ["Nivel AIS", result.ais_label, "—"],
            ],
            [9.0, 3.5, 2.5],
        )
        for desc in AIS_DESCRIPTIONS:
            if result.ais_label.startswith(desc["level"].split("—")[0].strip()):
                _heading("Interpretación clínica", 2)
                for label, key in [
                    ("Descripción clínica", "clinical"), ("Ejemplos", "examples"),
                    ("Pronóstico", "prognosis"), ("Contexto", "context"),
                ]:
                    _para(f"{label}: {_strip_html(desc[key])}")
                break

    # 4. Limitations
    _heading("4. Alcance y Limitaciones", 1)
    for item in [
        "Se desprecia la resistencia aerodinámica.",
        "Se asume rodadura pura sin deslizamiento.",
        "Los μ_r son valores típicos y varían con condiciones reales.",
        "La curva de Rosén & Sander aplica a condiciones europeas; su extrapolación requiere precaución.",
        "ESTE DOCUMENTO NO DEBE USARSE PARA ANÁLISIS FORENSE.",
    ]:
        _para(f"• {item}")

    # 5. References
    _heading("5. Referencias (APA 7.ª ed.)", 1)
    for ref in APA_REFERENCES:
        _para(_strip_html(ref))

    _para(
        f"\nGenerado automáticamente — Simulador de Impacto · "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}",
        color=RGBColor(0x99, 0x99, 0x99),
    )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
